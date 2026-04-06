#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# tabs/leaves_tab.py

import os
import logging
from datetime import datetime, date, timedelta

from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QComboBox, QPushButton,
    QMessageBox, QDialog, QFormLayout, QDateEdit, QTextEdit, QSpinBox,
    QDoubleSpinBox, QDialogButtonBox, QTabWidget, QInputDialog, QApplication,
    QLineEdit, QGroupBox, QGridLayout, QAbstractItemView, QCompleter,
    QFileDialog
)
from PyQt5.QtCore import Qt, QDate, QUrl
from PyQt5.QtGui import QFont

from database import DatabaseManager
from utils import make_table, fill_table, btn, can_add, can_edit, can_delete, can_approve
from constants import BTN_SUCCESS, BTN_PRIMARY, BTN_DANGER, BTN_GRAY, BTN_WARNING, BTN_PURPLE

logger = logging.getLogger(__name__)


# ============================================================
# دوال حساب الإجازات — المصدر الموحد الوحيد
# ============================================================
# هذه الدوال على مستوى الـ module لتجنب التكرار بين
# LeavesTab و LeaveRequestDialog اللتين كانتا تحتويان
# على نسخ مطابقة من نفس الكود.
# ============================================================

def _overlap_days(start_str: str, end_str: str, year: int) -> float:
    """
    حساب عدد أيام التداخل بين إجازة وسنة معينة.
    مثال: إجازة من 28/12/2024 إلى 3/1/2025 → تداخل مع 2025 = 3 أيام.
    """
    try:
        s          = date.fromisoformat(start_str)
        e          = date.fromisoformat(end_str)
        year_start = date(year, 1, 1)
        year_end   = date(year, 12, 31)
        ov_start   = max(s, year_start)
        ov_end     = min(e, year_end)
        if ov_start <= ov_end:
            return float((ov_end - ov_start).days + 1)
    except (ValueError, TypeError):
        pass
    return 0.0


def calc_used_days(db: DatabaseManager, emp_id: int,
                   leave_type_id: int, year: int) -> float:
    """
    الأيام المستخدمة (موافق) لنوع إجازة معين في سنة معينة.
    يحسب التداخل الدقيق للإجازات التي تمتد عبر نهاية/بداية السنة.
    """
    rows = db.fetch_all(
        """SELECT start_date, end_date FROM leave_requests
           WHERE employee_id   = ?
             AND leave_type_id = ?
             AND status        = 'موافق'
             AND start_date   <= ?
             AND end_date     >= ?""",
        (emp_id, leave_type_id,
         date(year, 12, 31).isoformat(),
         date(year, 1,  1).isoformat())
    )
    return sum(_overlap_days(s, e, year) for s, e in rows)


def calc_pending_days(db: DatabaseManager, emp_id: int,
                      leave_type_id: int, year: int) -> float:
    """
    الأيام المعلقة (قيد المراجعة) لنوع إجازة معين في سنة معينة.
    """
    rows = db.fetch_all(
        """SELECT start_date, end_date FROM leave_requests
           WHERE employee_id   = ?
             AND leave_type_id = ?
             AND status        = 'قيد المراجعة'
             AND start_date   <= ?
             AND end_date     >= ?""",
        (emp_id, leave_type_id,
         date(year, 12, 31).isoformat(),
         date(year, 1,  1).isoformat())
    )
    return sum(_overlap_days(s, e, year) for s, e in rows)


def calc_entitlement(db: DatabaseManager, emp_id: int,
                     leave_type_id: int, year: int) -> float:
    """
    المستحق من إجازة معينة لموظف في سنة معينة.

    - الإجازات السنوية (is_annual=1): تُحسَب من قواعد الخدمة + العمر.
    - الإجازات الأخرى: تُقرأ من leave_balance.
    """
    lt = db.fetch_one(
        "SELECT is_annual FROM leave_types WHERE id = ?", (leave_type_id,))
    if not lt:
        return 0.0

    if lt[0]:  # سنوية
        emp = db.fetch_one(
            "SELECT hire_date, birth_date FROM employees WHERE id = ?", (emp_id,))
        if not emp or not emp[0]:
            return 0.0

        year_start = date(year, 1, 1)
        try:
            hire_date = date.fromisoformat(emp[0])
        except ValueError:
            return 0.0

        service_years = max(0.0, (year_start - hire_date).days / 365.25) \
            if hire_date <= year_start else 0.0

        age = 0
        if emp[1]:
            try:
                birth = date.fromisoformat(emp[1])
                age   = year_start.year - birth.year
                if (year_start.month, year_start.day) < (birth.month, birth.day):
                    age -= 1
            except ValueError:
                pass

        rules = db.fetch_all(
            "SELECT from_year, to_year, days FROM leave_rules ORDER BY from_year")
        base_days = 0
        for from_y, to_y, days in rules:
            if from_y <= service_years <= to_y:
                base_days = days
                break

        if 0 < age <= 18 or age >= 50:
            return float(max(base_days, 20))
        return float(base_days)

    else:  # غير سنوية — من leave_balance
        bal = db.fetch_one(
            """SELECT total_days FROM leave_balance
               WHERE employee_id = ? AND leave_type_id = ? AND year = ?""",
            (emp_id, leave_type_id, year))
        return float(bal[0]) if bal else 0.0


# ============================================================
# الكلاس الرئيسي
# ============================================================
class LeavesTab(QWidget):
    def __init__(self, db: DatabaseManager, user: dict, comm=None):
        super().__init__()
        self.db   = db
        self.user = user
        self.comm = comm
        self._build()
        self._load_requests()
        self._load_balance()
        self._load_types()
        if self.comm:
            self.comm.dataChanged.connect(self._on_data_changed)

    def _on_data_changed(self, data_type: str, data):
        if data_type == 'employee':
            self._refresh_filters()
        elif data_type == 'leave_type':
            self._load_types()
            self._refresh_filters()
        elif data_type == 'leave_request':
            self._load_requests()
        elif data_type == 'leave_balance':
            self._load_balance()

    def _build(self):
        layout = QVBoxLayout(self)
        tabs   = QTabWidget()
        tabs.addTab(self._build_requests_tab(), "📋 طلبات الإجازة")
        tabs.addTab(self._build_balance_tab(),  "📊 أرصدة الإجازات")
        tabs.addTab(self._build_types_tab(),    "⚙️ أنواع الإجازات")
        layout.addWidget(tabs)
        self.inner_tabs = tabs

    # ----------------------------------------------------------
    # تبويب طلبات الإجازة
    # ----------------------------------------------------------
    def _build_requests_tab(self):
        w   = QWidget()
        lay = QVBoxLayout(w)

        # --- شريط الأزرار ---
        tools = QHBoxLayout()
        self.btn_new     = btn("➕ طلب جديد", BTN_SUCCESS, self._new_request)
        self.btn_edit    = btn("✏️ تعديل",    BTN_PRIMARY, self._edit_request)
        self.btn_approve = btn("✅ موافقة",   BTN_PRIMARY, self._approve)
        self.btn_reject  = btn("❌ رفض",      BTN_DANGER,  self._reject)
        self.btn_refresh = btn("🔄 تحديث",   BTN_GRAY,    self._load_requests)
        self.btn_export  = btn("📊 Excel",   BTN_PURPLE,  self._export_requests_excel)
        for b in (self.btn_new, self.btn_edit, self.btn_approve,
                  self.btn_reject, self.btn_refresh, self.btn_export):
            tools.addWidget(b)
        tools.addStretch()
        lay.addLayout(tools)

        # --- شريط الفلاتر ---
        filters = QHBoxLayout()

        filters.addWidget(QLabel("الحالة:"))
        self.req_status_filter = QComboBox()
        self.req_status_filter.addItems(
            ["جميع الطلبات", "قيد المراجعة", "موافق", "مرفوض"])
        self.req_status_filter.currentIndexChanged.connect(self._load_requests)
        filters.addWidget(self.req_status_filter)

        # ← فلتر نوع الإجازة (جديد)
        filters.addWidget(QLabel("النوع:"))
        self.req_type_filter = QComboBox()
        self.req_type_filter.setMinimumWidth(140)
        self._load_type_combo(self.req_type_filter)
        self.req_type_filter.currentIndexChanged.connect(self._load_requests)
        filters.addWidget(self.req_type_filter)

        filters.addWidget(QLabel("الموظف:"))
        self.req_emp_filter = QComboBox()
        self.req_emp_filter.setMinimumWidth(200)
        self._load_emp_combo(self.req_emp_filter)
        self.req_emp_filter.setEditable(True)
        self.req_emp_filter.completer().setCompletionMode(QCompleter.PopupCompletion)
        self.req_emp_filter.completer().setFilterMode(Qt.MatchContains)
        self.req_emp_filter.currentIndexChanged.connect(self._load_requests)
        filters.addWidget(self.req_emp_filter)

        filters.addStretch()
        lay.addLayout(filters)

        # --- الجدول ---
        self.req_table = make_table([
            "id", "الموظف", "نوع الإجازة", "من", "إلى",
            "الأيام", "السبب", "الحالة", "تاريخ الطلب"])
        self.req_table.setColumnHidden(0, True)
        self.req_table.setSortingEnabled(False)
        lay.addWidget(self.req_table)

        self.balance_summary = QLabel("")
        self.balance_summary.setStyleSheet(
            "font-weight:bold; padding:8px; background:#e3f2fd; border-radius:4px;")
        self.balance_summary.hide()
        lay.addWidget(self.balance_summary)

        self._apply_requests_permissions()
        return w

    def _apply_requests_permissions(self):
        role = self.user['role']
        self.btn_new.setVisible(can_add(role))
        self.btn_edit.setVisible(can_edit(role))
        self.btn_approve.setVisible(can_approve(role))
        self.btn_reject.setVisible(can_approve(role))

    # ----------------------------------------------------------
    # تبويب أرصدة الإجازات
    # ----------------------------------------------------------
    def _build_balance_tab(self):
        w   = QWidget()
        lay = QVBoxLayout(w)

        tools = QHBoxLayout()
        self.bal_year = QSpinBox()
        self.bal_year.setRange(2020, 2050)
        self.bal_year.setValue(date.today().year)
        tools.addWidget(QLabel("السنة:"))
        tools.addWidget(self.bal_year)

        self.bal_emp_filter = QComboBox()
        self.bal_emp_filter.setMinimumWidth(200)
        self._load_emp_combo(self.bal_emp_filter)
        self.bal_emp_filter.setEditable(True)
        self.bal_emp_filter.completer().setCompletionMode(QCompleter.PopupCompletion)
        self.bal_emp_filter.completer().setFilterMode(Qt.MatchContains)
        tools.addWidget(QLabel("الموظف:"))
        tools.addWidget(self.bal_emp_filter)

        self.bal_type_filter = QComboBox()
        self.bal_type_filter.setMinimumWidth(140)
        self._load_type_combo(self.bal_type_filter)
        self.bal_type_filter.setEditable(True)
        self.bal_type_filter.completer().setCompletionMode(QCompleter.PopupCompletion)
        self.bal_type_filter.completer().setFilterMode(Qt.MatchContains)
        tools.addWidget(QLabel("النوع:"))
        tools.addWidget(self.bal_type_filter)

        self.btn_load_balance   = btn("عرض",          BTN_PRIMARY, self._load_balance)
        self.btn_add_balance    = btn("➕ إضافة رصيد", BTN_SUCCESS, self._add_balance)
        self.btn_edit_balance   = btn("✏️ تعديل",      BTN_PRIMARY, self._edit_balance)
        self.btn_delete_balance = btn("🗑️ حذف",        BTN_DANGER,  self._delete_balance)
        for b in (self.btn_load_balance, self.btn_add_balance,
                  self.btn_edit_balance, self.btn_delete_balance):
            tools.addWidget(b)
        tools.addStretch()
        lay.addLayout(tools)

        self.bal_table = make_table([
            "balance_id", "الموظف", "نوع الإجازة", "السنة",
            "المستحق", "المستخدم", "المعلق", "المتبقي"])
        self.bal_table.setColumnHidden(0, True)
        self.bal_table.setSortingEnabled(False)
        self.bal_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.bal_table.setSelectionMode(QAbstractItemView.SingleSelection)
        lay.addWidget(self.bal_table)

        self._apply_balance_permissions()
        return w

    def _apply_balance_permissions(self):
        role = self.user['role']
        self.btn_add_balance.setVisible(can_add(role))
        self.btn_edit_balance.setVisible(can_edit(role))
        self.btn_delete_balance.setVisible(can_delete(role))

    # ----------------------------------------------------------
    # تبويب أنواع الإجازات
    # ----------------------------------------------------------
    def _build_types_tab(self):
        w   = QWidget()
        lay = QVBoxLayout(w)

        tools = QHBoxLayout()
        self.btn_new_type    = btn("➕ نوع جديد", BTN_SUCCESS, self._new_type)
        self.btn_edit_type   = btn("✏️ تعديل",   BTN_PRIMARY, self._edit_type)
        self.btn_delete_type = btn("🗑️ حذف",     BTN_DANGER,  self._delete_type)
        for b in (self.btn_new_type, self.btn_edit_type, self.btn_delete_type):
            tools.addWidget(b)
        tools.addStretch()
        lay.addLayout(tools)

        self.types_table = make_table([
            "#", "النوع", "أيام/سنة", "مدفوعة", "ترحيل",
            "سنوي", "الحد الأقصى للطلبات"])
        lay.addWidget(self.types_table)
        self._load_types()

        self._apply_types_permissions()
        return w

    def _apply_types_permissions(self):
        role = self.user['role']
        self.btn_new_type.setVisible(can_add(role))
        self.btn_edit_type.setVisible(can_edit(role))
        self.btn_delete_type.setVisible(can_delete(role))

    # ----------------------------------------------------------
    # دوال مساعدة مشتركة
    # ----------------------------------------------------------
    def _load_emp_combo(self, combo: QComboBox):
        """تحميل قائمة الموظفين النشطين في combobox."""
        combo.clear()
        combo.addItem("جميع الموظفين", None)
        for eid, name in self.db.fetch_all(
                "SELECT id, first_name||' '||last_name "
                "FROM employees WHERE status='نشط' ORDER BY first_name"):
            combo.addItem(name, eid)

    def _load_type_combo(self, combo: QComboBox):
        """تحميل قائمة أنواع الإجازات في combobox."""
        combo.clear()
        combo.addItem("جميع الأنواع", None)
        for tid, tname in self.db.fetch_all(
                "SELECT id, name FROM leave_types ORDER BY name"):
            combo.addItem(tname, tid)

    def _refresh_filters(self):
        """إعادة تحميل جميع قوائم الفلاتر مع الحفاظ على الاختيار الحالي."""
        for combo, loader in [
            (self.req_emp_filter,  self._load_emp_combo),
            (self.bal_emp_filter,  self._load_emp_combo),
            (self.req_type_filter, self._load_type_combo),
            (self.bal_type_filter, self._load_type_combo),
        ]:
            prev = combo.currentData()
            loader(combo)
            idx = combo.findData(prev)
            if idx >= 0:
                combo.setCurrentIndex(idx)
            combo.setEditable(True)
            combo.completer().setCompletionMode(QCompleter.PopupCompletion)
            combo.completer().setFilterMode(Qt.MatchContains)

    # ----------------------------------------------------------
    # عمليات طلبات الإجازة
    # ----------------------------------------------------------
    def _load_requests(self):
        status_filter = self.req_status_filter.currentText()
        emp_id        = self.req_emp_filter.currentData()
        type_id       = self.req_type_filter.currentData()   # ← جديد

        q = """
            SELECT lr.id,
                   e.first_name || ' ' || e.last_name,
                   lt.name,
                   lr.start_date, lr.end_date, lr.days_count,
                   lr.reason, lr.status, lr.created_at
            FROM leave_requests lr
            JOIN employees   e  ON lr.employee_id   = e.id
            JOIN leave_types lt ON lr.leave_type_id = lt.id
            WHERE 1=1
        """
        params = []
        if status_filter != "جميع الطلبات":
            q += " AND lr.status = ?"
            params.append(status_filter)
        if emp_id is not None:
            q += " AND lr.employee_id = ?"
            params.append(emp_id)
        if type_id is not None:                              # ← جديد
            q += " AND lr.leave_type_id = ?"
            params.append(type_id)
        q += " ORDER BY lr.created_at DESC"

        data = self.db.fetch_all(q, params if params else None)

        if not data:
            self.req_table.setRowCount(0)
            self.balance_summary.hide()
            return

        fill_table(
            self.req_table,
            [[row[0]] + list(row[1:]) for row in data],
            colors={7: lambda v: (
                "#388E3C" if v == "موافق"
                else "#D32F2F" if v == "مرفوض"
                else "#F57C00"
            )}
        )

        if emp_id:
            self._show_balance_summary(emp_id)
        else:
            self.balance_summary.hide()

    def _show_balance_summary(self, emp_id: int):
        year  = date.today().year
        lines = []
        for tid, tname in self.db.fetch_all("SELECT id, name FROM leave_types"):
            entitled  = calc_entitlement(self.db, emp_id, tid, year)
            used      = calc_used_days(self.db, emp_id, tid, year)
            pending   = calc_pending_days(self.db, emp_id, tid, year)
            remaining = entitled - used - pending
            lines.append(
                f"{tname}: المستحق {entitled:.0f} | "
                f"المستخدم {used:.0f} | المعلق {pending:.0f} | "
                f"المتبقي {remaining:.0f}"
            )
        self.balance_summary.setText("\n".join(lines))
        self.balance_summary.show()

    def _new_request(self):
        dlg = LeaveRequestDialog(self, self.db, self.user, None, self.comm)
        if dlg.exec_() == QDialog.Accepted:
            self._load_requests()
            if self.comm:
                self.comm.dataChanged.emit('leave_request', {'action': 'add'})

    def _edit_request(self):
        row = self.req_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "خطأ", "اختر طلباً أولاً")
            return
        req_id = int(self.req_table.item(row, 0).text())
        dlg    = LeaveRequestDialog(self, self.db, self.user, req_id, self.comm)
        if dlg.exec_() == QDialog.Accepted:
            self._load_requests()
            if self.comm:
                self.comm.dataChanged.emit('leave_request', {'action': 'edit', 'id': req_id})

    def _approve(self):
        row = self.req_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "خطأ", "اختر طلباً أولاً")
            return
        req_id = int(self.req_table.item(row, 0).text())

        cur = self.db.fetch_one(
            "SELECT status FROM leave_requests WHERE id = ?", (req_id,))
        if not cur:
            QMessageBox.critical(self, "خطأ", "لم يتم العثور على الطلب")
            return
        if cur[0] == 'موافق':
            QMessageBox.warning(self, "خطأ", "هذا الطلب موافق عليه مسبقاً")
            return

        req = self.db.fetch_one(
            "SELECT employee_id, leave_type_id, start_date, end_date, days_count "
            "FROM leave_requests WHERE id = ?", (req_id,))
        if not req:
            return
        emp_id, leave_type_id, start_date, end_date, days = req

        # التحقق من الرصيد للإجازات السنوية المدفوعة
        lt = self.db.fetch_one(
            "SELECT is_annual, paid FROM leave_types WHERE id = ?", (leave_type_id,))
        if lt and lt[0] == 1 and lt[1] == 1:
            year      = date.today().year
            entitled  = calc_entitlement(self.db, emp_id, leave_type_id, year)
            used      = calc_used_days(self.db, emp_id, leave_type_id, year)
            pending   = calc_pending_days(self.db, emp_id, leave_type_id, year)
            remaining = entitled - used - pending
            if days > remaining:
                QMessageBox.warning(
                    self, "خطأ",
                    f"الرصيد المتبقي غير كافٍ.\nالمطلوب: {days} | المتاح: {remaining:.0f}")
                return

        self.db.execute_query(
            "UPDATE leave_requests SET status='موافق', approved_by=?, approved_at=? WHERE id=?",
            (self.user['id'], datetime.now().isoformat(), req_id))

        # تحديث leave_balance للإجازات غير السنوية
        if lt and lt[0] == 0:
            year    = date.today().year
            balance = self.db.fetch_one(
                """SELECT id FROM leave_balance
                   WHERE employee_id = ? AND leave_type_id = ? AND year = ?""",
                (emp_id, leave_type_id, year))
            if balance:
                self.db.execute_query(
                    """UPDATE leave_balance
                       SET used_days = used_days + ?, pending_days = pending_days - ?
                       WHERE id = ?""",
                    (days, days, balance[0]))
            else:
                self.db.execute_query(
                    """INSERT INTO leave_balance
                       (employee_id, leave_type_id, year, total_days, used_days, pending_days)
                       VALUES (?, ?, ?, 0, ?, 0)""",
                    (emp_id, leave_type_id, year, days))

        self._add_leave_attendance(emp_id, start_date, end_date)
        self.db.log_action("موافقة على إجازة", "leave_requests", req_id,
                           None, {"employee_id": emp_id,
                                  "start_date": start_date, "end_date": end_date})

        if self.comm:
            self.comm.dataChanged.emit('leave_request', {'action': 'approve', 'id': req_id})
            self.comm.dataChanged.emit('attendance',
                                       {'action': 'update', 'employee_id': emp_id,
                                        'start': start_date, 'end': end_date})

        self._load_requests()
        QMessageBox.information(self, "نجاح",
                                "تمت الموافقة على الطلب وتحديث سجلات الحضور")

    def _reject(self):
        row = self.req_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "خطأ", "اختر طلباً أولاً")
            return
        req_id = int(self.req_table.item(row, 0).text())

        req = self.db.fetch_one(
            "SELECT employee_id, start_date, end_date, status, leave_type_id, days_count "
            "FROM leave_requests WHERE id = ?", (req_id,))
        if not req:
            QMessageBox.critical(self, "خطأ", "لم يتم العثور على الطلب")
            return
        emp_id, start_date, end_date, old_status, leave_type_id, days = req

        reason, ok = QInputDialog.getMultiLineText(
            self, "سبب الرفض", "الرجاء كتابة سبب رفض الطلب:")
        if not ok:
            return

        self.db.execute_query(
            "UPDATE leave_requests SET status='مرفوض', approved_by=?, approved_at=?, notes=? WHERE id=?",
            (self.user['id'], datetime.now().isoformat(), reason, req_id))

        if old_status == 'موافق':
            self._remove_leave_attendance(emp_id, start_date, end_date)
            lt = self.db.fetch_one(
                "SELECT is_annual FROM leave_types WHERE id = ?", (leave_type_id,))
            if lt and lt[0] == 0:
                year    = date.today().year
                balance = self.db.fetch_one(
                    "SELECT id FROM leave_balance "
                    "WHERE employee_id=? AND leave_type_id=? AND year=?",
                    (emp_id, leave_type_id, year))
                if balance:
                    self.db.execute_query(
                        """UPDATE leave_balance
                           SET used_days = used_days - ?, pending_days = pending_days + ?
                           WHERE id = ?""",
                        (days, days, balance[0]))

        self.db.log_action("رفض إجازة", "leave_requests", req_id,
                           None, {"employee_id": emp_id, "reason": reason})

        if self.comm:
            self.comm.dataChanged.emit('leave_request', {'action': 'reject', 'id': req_id})
            if old_status == 'موافق':
                self.comm.dataChanged.emit('attendance',
                                           {'action': 'update', 'employee_id': emp_id,
                                            'start': start_date, 'end': end_date})

        self._load_requests()

    def _export_requests_excel(self):
        try:
            import pandas as pd
        except ImportError:
            QMessageBox.critical(self, "خطأ", "pip install pandas openpyxl")
            return

        status_filter = self.req_status_filter.currentText()
        emp_id        = self.req_emp_filter.currentData()
        type_id       = self.req_type_filter.currentData()

        q = """
            SELECT e.employee_code, e.first_name||' '||e.last_name, lt.name,
                   lr.start_date, lr.end_date, lr.days_count,
                   lr.reason, lr.status, lr.created_at
            FROM leave_requests lr
            JOIN employees   e  ON lr.employee_id   = e.id
            JOIN leave_types lt ON lr.leave_type_id = lt.id
            WHERE 1=1
        """
        params = []
        if status_filter != "جميع الطلبات":
            q += " AND lr.status = ?"
            params.append(status_filter)
        if emp_id is not None:
            q += " AND lr.employee_id = ?"
            params.append(emp_id)
        if type_id is not None:
            q += " AND lr.leave_type_id = ?"
            params.append(type_id)
        q += " ORDER BY lr.created_at DESC"

        data = self.db.fetch_all(q, params if params else None)
        if not data:
            QMessageBox.warning(self, "تنبيه", "لا توجد بيانات للتصدير")
            return

        df   = pd.DataFrame(data, columns=[
            "الرقم الوظيفي", "الموظف", "نوع الإجازة",
            "من", "إلى", "الأيام", "السبب", "الحالة", "تاريخ الطلب"])
        path, _ = QFileDialog.getSaveFileName(
            self, "حفظ التقرير", "leave_requests.xlsx", "Excel (*.xlsx)")
        if path:
            df.to_excel(path, index=False)
            QMessageBox.information(self, "نجاح", f"تم تصدير {len(df)} سجل")

    # ----------------------------------------------------------
    # سجلات الحضور المرتبطة بالإجازات
    # ----------------------------------------------------------
    def _add_leave_attendance(self, emp_id: int, start_date: str, end_date: str):
        s, e    = date.fromisoformat(start_date), date.fromisoformat(end_date)
        current = s
        while current <= e:
            day_str = current.isoformat()
            self.db.execute_query(
                "DELETE FROM attendance WHERE employee_id=? AND punch_date=?",
                (emp_id, day_str))
            self.db.execute_query(
                """INSERT INTO attendance
                   (employee_id, punch_date, status,
                    work_hours, overtime_hours, late_minutes, early_leave_minutes)
                   VALUES (?,?,?,?,?,?,?)""",
                (emp_id, day_str, "إجازة", 0, 0, 0, 0))
            current += timedelta(days=1)

    def _remove_leave_attendance(self, emp_id: int,
                                  start_date: str, end_date: str):
        self.db.execute_query(
            """DELETE FROM attendance
               WHERE employee_id=? AND punch_date BETWEEN ? AND ? AND status='إجازة'""",
            (emp_id, start_date, end_date))

    # ----------------------------------------------------------
    # عمليات الأرصدة
    # ----------------------------------------------------------
    def _load_balance(self):
        year           = self.bal_year.value()
        emp_id         = self.bal_emp_filter.currentData()
        type_id_filter = self.bal_type_filter.currentData()

        emp_params = []
        emp_cond   = ""
        if emp_id is not None:
            emp_cond   = " AND e.id = ?"
            emp_params = [emp_id]

        employees = self.db.fetch_all(
            f"""SELECT e.id, e.first_name||' '||e.last_name,
                       e.hire_date, e.birth_date
                FROM employees e
                WHERE e.status = 'نشط' {emp_cond}""",
            emp_params if emp_params else None)

        leave_types = self.db.fetch_all(
            "SELECT id, name, is_annual FROM leave_types")

        # جلب الأرصدة المحفوظة لهذه السنة دفعة واحدة
        stored = {}
        for r in self.db.fetch_all(
                "SELECT id, employee_id, leave_type_id, total_days, used_days, pending_days "
                "FROM leave_balance WHERE year = ?", (year,)):
            stored[(r[1], r[2])] = {
                'id': r[0], 'total': r[3], 'used': r[4], 'pending': r[5]}

        # جلب كل الطلبات المتقاطعة مع السنة دفعة واحدة
        y_start = date(year, 1, 1).isoformat()
        y_end   = date(year, 12, 31).isoformat()
        used_map    = {}
        pending_map = {}
        for r in self.db.fetch_all(
                """SELECT employee_id, leave_type_id, start_date, end_date, status
                   FROM leave_requests
                   WHERE status IN ('موافق', 'قيد المراجعة')
                     AND start_date <= ? AND end_date >= ?""",
                (y_end, y_start)):
            key  = (r[0], r[1])
            days = _overlap_days(r[2], r[3], year)
            if r[4] == 'موافق':
                used_map[key]    = used_map.get(key, 0.0) + days
            else:
                pending_map[key] = pending_map.get(key, 0.0) + days

        display = []
        for emp in employees:
            emp_id_r, emp_name, hire_str, birth_str = emp
            for lt in leave_types:
                type_id, type_name, is_annual = lt
                if type_id_filter is not None and type_id != type_id_filter:
                    continue
                key = (emp_id_r, type_id)

                if is_annual:
                    entitled   = calc_entitlement(self.db, emp_id_r, type_id, year)
                    used       = used_map.get(key, 0.0)
                    pending    = pending_map.get(key, 0.0)
                    remaining  = entitled - used - pending
                    balance_id = None
                else:
                    bd         = stored.get(key, {'total': 0, 'used': 0, 'pending': 0})
                    entitled   = bd['total']
                    used       = bd['used']
                    pending    = bd['pending']
                    remaining  = entitled - used - pending
                    balance_id = bd.get('id')

                display.append((
                    balance_id, emp_name, type_name, year,
                    f"{entitled:.0f}", f"{used:.0f}",
                    f"{pending:.0f}", f"{remaining:.0f}"
                ))

        fill_table(self.bal_table, display, colors={
            7: lambda v: "#D32F2F" if float(v) < 0 else "#388E3C" if float(v) > 0 else None
        })

    def _add_balance(self):
        dlg = AddBalanceDialog(self, self.db, self.user, self.bal_year.value(), self.comm)
        if dlg.exec_() == QDialog.Accepted:
            self._load_balance()
            if self.comm:
                self.comm.dataChanged.emit('leave_balance', {'action': 'add'})

    def _edit_balance(self):
        row = self.bal_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "خطأ", "اختر رصيداً أولاً")
            return
        bal_id = self.bal_table.item(row, 0).text()
        if not bal_id or bal_id == "None":
            QMessageBox.warning(self, "خطأ", "لا يمكن تعديل هذا الرصيد (سنوي)")
            return
        bal_id  = int(bal_id)
        bal     = self.db.fetch_one("SELECT * FROM leave_balance WHERE id=?", (bal_id,))
        if not bal:
            QMessageBox.critical(self, "خطأ", "لم يتم العثور على الرصيد")
            return
        dlg = AddBalanceDialog(self, self.db, self.user, bal[3], self.comm, bal)
        if dlg.exec_() == QDialog.Accepted:
            self._load_balance()
            if self.comm:
                self.comm.dataChanged.emit('leave_balance', {'action': 'edit', 'id': bal_id})

    def _delete_balance(self):
        row = self.bal_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "خطأ", "اختر رصيداً أولاً")
            return
        bal_id = self.bal_table.item(row, 0).text()
        if not bal_id or bal_id == "None":
            QMessageBox.warning(self, "خطأ", "لا يمكن حذف هذا الرصيد (سنوي)")
            return
        bal_id   = int(bal_id)
        bal_data = self.db.fetch_one("SELECT * FROM leave_balance WHERE id=?", (bal_id,))
        if not bal_data:
            QMessageBox.critical(self, "خطأ", "لم يتم العثور على الرصيد")
            return
        if QMessageBox.question(
                self, "تأكيد", "هل أنت متأكد من حذف هذا الرصيد؟",
                QMessageBox.Yes | QMessageBox.No) == QMessageBox.Yes:
            self.db.execute_query("DELETE FROM leave_balance WHERE id=?", (bal_id,))
            self.db.log_action("حذف رصيد إجازة", "leave_balance", bal_id,
                               {"employee_id": bal_data[1], "total_days": bal_data[4]}, None)
            self._load_balance()
            if self.comm:
                self.comm.dataChanged.emit('leave_balance', {'action': 'delete', 'id': bal_id})

    # ----------------------------------------------------------
    # عمليات أنواع الإجازات
    # ----------------------------------------------------------
    def _load_types(self):
        data = self.db.fetch_all(
            "SELECT id, name, days_per_year, paid, carry_over, is_annual, max_requests "
            "FROM leave_types ORDER BY name")
        rows = [
            (r[0], r[1], r[2],
             "نعم" if r[3] else "لا",
             "نعم" if r[4] else "لا",
             "نعم" if r[5] else "لا",
             r[6] or "∞")
            for r in data
        ]
        fill_table(self.types_table, rows)

    def _new_type(self):
        dlg = EditLeaveTypeDialog(self, self.db, self.user, self.comm)
        if dlg.exec_() == QDialog.Accepted:
            self._load_types()
            self._refresh_filters()
            if self.comm:
                self.comm.dataChanged.emit('leave_type', {'action': 'add'})

    def _edit_type(self):
        row = self.types_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "خطأ", "اختر نوعاً أولاً")
            return
        type_id = int(self.types_table.item(row, 0).text())
        dlg     = EditLeaveTypeDialog(self, self.db, self.user, self.comm, type_id)
        if dlg.exec_() == QDialog.Accepted:
            self._load_types()
            self._refresh_filters()
            if self.comm:
                self.comm.dataChanged.emit('leave_type', {'action': 'edit', 'id': type_id})

    def _delete_type(self):
        row = self.types_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "خطأ", "اختر نوعاً أولاً")
            return
        type_id   = int(self.types_table.item(row, 0).text())
        type_data = self.db.fetch_one(
            "SELECT name, days_per_year FROM leave_types WHERE id=?", (type_id,))
        if not type_data:
            QMessageBox.critical(self, "خطأ", "لم يتم العثور على النوع")
            return

        req_cnt = self.db.fetch_one(
            "SELECT COUNT(*) FROM leave_requests WHERE leave_type_id=?", (type_id,))[0]
        bal_cnt = self.db.fetch_one(
            "SELECT COUNT(*) FROM leave_balance WHERE leave_type_id=?",  (type_id,))[0]

        if req_cnt > 0 or bal_cnt > 0:
            msg = "هذا النوع مستخدم في:\n"
            if req_cnt: msg += f"- {req_cnt} طلب إجازة\n"
            if bal_cnt: msg += f"- {bal_cnt} رصيد إجازة\n"
            msg += "سيتم حذف جميع السجلات المرتبطة. هل أنت متأكد؟"
            if QMessageBox.question(
                    self, "تأكيد", msg,
                    QMessageBox.Yes | QMessageBox.No) == QMessageBox.No:
                return
            self.db.execute_query(
                "DELETE FROM leave_requests WHERE leave_type_id=?", (type_id,))
            self.db.execute_query(
                "DELETE FROM leave_balance WHERE leave_type_id=?",  (type_id,))

        if QMessageBox.question(
                self, "تأكيد", "هل أنت متأكد من حذف هذا النوع؟",
                QMessageBox.Yes | QMessageBox.No) == QMessageBox.Yes:
            if self.db.execute_query("DELETE FROM leave_types WHERE id=?", (type_id,)):
                self.db.log_action("حذف نوع إجازة", "leave_types", type_id,
                                   {"name": type_data[0]}, None)
                self._load_types()
                self._refresh_filters()
                if self.comm:
                    self.comm.dataChanged.emit('leave_type', {'action': 'delete', 'id': type_id})
                QMessageBox.information(self, "نجاح", "تم حذف نوع الإجازة بنجاح")
            else:
                QMessageBox.critical(self, "خطأ", "فشل في حذف نوع الإجازة")


# ============================================================
# نافذة طلب إجازة جديد / تعديل
# ============================================================
class LeaveRequestDialog(QDialog):
    def __init__(self, parent, db: DatabaseManager, user: dict,
                 request_id=None, comm=None):
        super().__init__(parent)
        self.db         = db
        self.user       = user
        self.comm       = comm
        self.request_id = request_id
        self.setWindowTitle("طلب إجازة جديد" if not request_id else "تعديل طلب إجازة")
        self.setFixedSize(500, 550)
        self.setLayoutDirection(Qt.RightToLeft)
        self._build()
        if request_id:
            self._load()

    def _build(self):
        lay  = QVBoxLayout(self)
        form = QFormLayout()
        form.setSpacing(10)

        self.emp = QComboBox()
        for eid, name in self.db.fetch_all(
                "SELECT id, first_name||' '||last_name "
                "FROM employees WHERE status='نشط'"):
            self.emp.addItem(name, eid)
        self.emp.setEditable(True)
        self.emp.completer().setCompletionMode(QCompleter.PopupCompletion)
        self.emp.completer().setFilterMode(Qt.MatchContains)

        self.ltype = QComboBox()
        for tid, tname, is_annual, max_req, paid in self.db.fetch_all(
                "SELECT id, name, is_annual, max_requests, paid FROM leave_types"):
            self.ltype.addItem(tname, (tid, is_annual, max_req, paid))
        self.ltype.setEditable(True)
        self.ltype.completer().setCompletionMode(QCompleter.PopupCompletion)
        self.ltype.completer().setFilterMode(Qt.MatchContains)

        self.start = QDateEdit()
        self.start.setCalendarPopup(True)
        self.start.setDate(QDate.currentDate())
        self.end = QDateEdit()
        self.end.setCalendarPopup(True)
        self.end.setDate(QDate.currentDate().addDays(1))
        self.start.dateChanged.connect(self._calc_days)
        self.end.dateChanged.connect(self._calc_days)

        self.days_lbl = QLabel("0 يوم")
        self.days_lbl.setStyleSheet("font-weight:bold; color:#1976D2;")

        self.balance_info = QLabel("")
        self.balance_info.setStyleSheet("color:#666; font-size:11px;")
        self.emp.currentIndexChanged.connect(self._update_balance_info)
        self.ltype.currentIndexChanged.connect(self._update_balance_info)

        self.reason = QTextEdit()
        self.reason.setMaximumHeight(80)

        form.addRow("الموظف:",     self.emp)
        form.addRow("نوع الإجازة:", self.ltype)
        form.addRow(self.balance_info)
        form.addRow("من:",          self.start)
        form.addRow("إلى:",         self.end)
        form.addRow("عدد الأيام:", self.days_lbl)
        form.addRow("السبب:",       self.reason)
        lay.addLayout(form)

        btn_box = QHBoxLayout()
        btn_box.addWidget(btn("حفظ وطباعة", BTN_SUCCESS, self._save_and_print))
        btn_box.addWidget(btn("حفظ فقط",    BTN_PRIMARY, self._save_only))
        btn_box.addWidget(btn("إلغاء",      BTN_DANGER,  self.reject))
        lay.addLayout(btn_box)

        self._calc_days()
        self._update_balance_info()

    def _calc_days(self):
        d1   = self.start.date().toPyDate()
        d2   = self.end.date().toPyDate()
        days = max(0, (d2 - d1).days + 1)
        self.days_lbl.setText(f"{days} يوم")

    def _update_balance_info(self):
        """عرض الرصيد المتبقي بجانب نوع الإجازة — يستخدم الدوال الموحدة."""
        emp_id    = self.emp.currentData()
        type_data = self.ltype.currentData()
        if not emp_id or not type_data:
            self.balance_info.setText("")
            return
        type_id   = type_data[0]
        year      = date.today().year
        entitled  = calc_entitlement(self.db, emp_id, type_id, year)
        used      = calc_used_days(self.db, emp_id, type_id, year)
        pending   = calc_pending_days(self.db, emp_id, type_id, year)
        remaining = entitled - used - pending
        self.balance_info.setText(
            f"الرصيد المتبقي: {remaining:.0f} يوم "
            f"(مستحق {entitled:.0f} | مستخدم {used:.0f} | معلق {pending:.0f})")

    def _validate(self):
        emp_id    = self.emp.currentData()
        type_data = self.ltype.currentData()
        if not type_data:
            QMessageBox.warning(self, "خطأ", "نوع الإجازة مطلوب")
            return False
        type_id, is_annual, max_requests, paid = type_data

        start_date = self.start.date().toString(Qt.ISODate)
        end_date   = self.end.date().toString(Qt.ISODate)

        # التحقق من التداخل مع طلبات أخرى
        exclude = self.request_id if self.request_id else 0
        row = self.db.fetch_one(
            """SELECT id FROM leave_requests
               WHERE employee_id = ? AND status IN ('قيد المراجعة', 'موافق')
                 AND id != ?
                 AND NOT (end_date < ? OR start_date > ?)""",
            (emp_id, exclude, start_date, end_date))
        if row:
            QMessageBox.warning(self, "خطأ",
                                "يوجد طلب إجازة آخر متداخل مع هذه الفترة.")
            return False

        # التحقق من الحد الأقصى للطلبات
        if max_requests and max_requests > 0:
            count = self.db.fetch_one(
                """SELECT COUNT(*) FROM leave_requests
                   WHERE employee_id=? AND leave_type_id=? AND status='موافق'""",
                (emp_id, type_id))[0]
            if count >= max_requests:
                QMessageBox.warning(
                    self, "خطأ",
                    f"هذا النوع يسمح بـ {max_requests} مرة فقط. استنفدت حصتك.")
                return False

        # التحقق من الرصيد للإجازات السنوية المدفوعة
        if is_annual and paid == 1:
            days      = int(self.days_lbl.text().split()[0])
            year      = date.today().year
            remaining = (calc_entitlement(self.db, emp_id, type_id, year)
                         - calc_used_days(self.db, emp_id, type_id, year)
                         - calc_pending_days(self.db, emp_id, type_id, year))
            if days > remaining:
                QMessageBox.warning(
                    self, "خطأ",
                    f"الرصيد المتبقي غير كافٍ. المتبقي: {remaining:.0f} يوم")
                return False

        return True

    def _save_request(self):
        emp_id    = self.emp.currentData()
        type_data = self.ltype.currentData()
        type_id   = type_data[0] if type_data else None
        start_d   = self.start.date().toString(Qt.ISODate)
        end_d     = self.end.date().toString(Qt.ISODate)
        days      = int(self.days_lbl.text().split()[0])
        reason    = self.reason.toPlainText().strip()

        if not emp_id or not type_id:
            QMessageBox.warning(self, "خطأ", "يرجى اختيار الموظف ونوع الإجازة")
            return None
        if days <= 0:
            QMessageBox.warning(self, "خطأ", "عدد الأيام يجب أن يكون أكبر من صفر")
            return None

        if self.request_id:
            old = self.db.fetch_one(
                "SELECT status, start_date, end_date FROM leave_requests WHERE id=?",
                (self.request_id,))
            if old and old[0] == 'موافق':
                self.db.execute_query(
                    """DELETE FROM attendance
                       WHERE employee_id=? AND punch_date BETWEEN ? AND ? AND status='إجازة'""",
                    (emp_id, old[1], old[2]))
            self.db.execute_query(
                """UPDATE leave_requests
                   SET employee_id=?, leave_type_id=?, start_date=?, end_date=?,
                       days_count=?, reason=?
                   WHERE id=?""",
                (emp_id, type_id, start_d, end_d, days, reason, self.request_id))
            return self.request_id
        else:
            ok = self.db.execute_query(
                """INSERT INTO leave_requests
                   (employee_id, leave_type_id, start_date, end_date,
                    days_count, reason, status, created_at)
                   VALUES (?,?,?,?,?,?,?,?)""",
                (emp_id, type_id, start_d, end_d, days, reason,
                 "قيد المراجعة", datetime.now().isoformat()))
            return self.db.last_id() if ok else None

    def _load(self):
        req = self.db.fetch_one(
            "SELECT * FROM leave_requests WHERE id=?", (self.request_id,))
        if not req:
            return
        idx = self.emp.findData(req[1])
        if idx >= 0:
            self.emp.setCurrentIndex(idx)
        for i in range(self.ltype.count()):
            if self.ltype.itemData(i)[0] == req[2]:
                self.ltype.setCurrentIndex(i)
                break
        if req[4]:
            self.start.setDate(QDate.fromString(str(req[4]), Qt.ISODate))
        if req[5]:
            self.end.setDate(QDate.fromString(str(req[5]), Qt.ISODate))
        self.reason.setText(req[6] or "")
        self._calc_days()
        self._update_balance_info()

    def _log_and_accept(self, req_id: int):
        self.db.log_action(
            "تعديل طلب إجازة" if self.request_id else "إضافة طلب إجازة",
            "leave_requests", req_id, None,
            {"employee_id":  self.emp.currentData(),
             "start_date":   self.start.date().toString(Qt.ISODate),
             "end_date":     self.end.date().toString(Qt.ISODate),
             "days":         int(self.days_lbl.text().split()[0])})
        self.accept()

    def _save_only(self):
        if not self._validate():
            return
        req_id = self._save_request()
        if req_id:
            QMessageBox.information(self, "تم", "تم حفظ الطلب بنجاح")
            self._log_and_accept(req_id)

    def _save_and_print(self):
        if not self._validate():
            return
        req_id = self._save_request()
        if req_id:
            self._print_request(req_id)
            self._log_and_accept(req_id)

    def _print_request(self, req_id: int):
        req = self.db.fetch_one(
            """SELECT lr.id, lr.employee_id, lr.leave_type_id,
                      lr.start_date, lr.end_date, lr.days_count,
                      lr.reason, lr.status, lr.approved_by, lr.approved_at,
                      lr.notes, lr.created_at,
                      e.first_name||' '||e.last_name,
                      e.national_id, d.name, e.employee_code, lt.name
               FROM leave_requests lr
               JOIN employees   e  ON lr.employee_id   = e.id
               LEFT JOIN departments d ON e.department_id = d.id
               JOIN leave_types lt ON lr.leave_type_id = lt.id
               WHERE lr.id = ?""", (req_id,))
        if not req:
            return

        company_name    = self.db.get_setting('company_name', 'الشركة')
        company_address = self.db.get_setting('company_address', '')
        company_phone   = self.db.get_setting('company_phone', '')
        logo_path       = self.db.get_setting('company_logo', '')

        emp_id   = req[1]
        type_id  = req[2]
        year     = date.today().year
        entitled = calc_entitlement(self.db, emp_id, type_id, year)
        used     = calc_used_days(self.db, emp_id, type_id, year)
        pending  = calc_pending_days(self.db, emp_id, type_id, year)
        remaining = entitled - used - pending

        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            QMessageBox.critical(
                self, "خطأ",
                "مكتبة python-docx غير مثبتة.\npip install python-docx")
            return

        doc      = Document()
        for section in doc.sections:
            for attr in ('top_margin', 'bottom_margin',
                         'left_margin', 'right_margin'):
                setattr(section, attr, Inches(0.5))

        # رأس الصفحة
        ht = doc.add_table(rows=1, cols=2)
        ht.autofit = False
        ht.columns[0].width = Inches(3)
        ht.columns[1].width = Inches(3)
        if logo_path and os.path.exists(logo_path):
            p = ht.cell(0, 0).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run().add_picture(logo_path, width=Inches(1.5))
        p = ht.cell(0, 1).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(company_name + "\n")
        r.bold = True
        r.font.size = Pt(14)
        if company_address:
            p.add_run(company_address + "\n").font.size = Pt(10)
        if company_phone:
            p.add_run("Tel: " + company_phone).font.size = Pt(10)

        doc.add_paragraph()
        t = doc.add_paragraph("İzin Talebi")
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t.runs[0].font.size = Pt(16)
        t.runs[0].bold = t.runs[0].underline = True
        doc.add_paragraph()

        table = doc.add_table(rows=10, cols=2)
        table.style     = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, value) in enumerate([
            ("Personel No",       req[15] or '---'),
            ("Ad Soyad",          req[12]),
            ("TC Kimlik",         req[13] or '---'),
            ("Departman",         req[14] or '---'),
            ("İzin Türü",         req[16]),
            ("Başlangıç Tarihi",  req[3]),
            ("Bitiş Tarihi",      req[4]),
            ("Gün Sayısı",        str(req[5])),
            ("Sebep",             req[6] or '---'),
            ("Talep Tarihi",      req[11]),
        ]):
            rc = table.rows[i].cells
            rc[0].text = label
            rc[1].text = value
            for cell in rc:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(11)

        doc.add_paragraph()
        bt = doc.add_paragraph("Yıllık İzin Bakiyesi")
        bt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bt.runs[0].font.size = Pt(14)
        bt.runs[0].bold = True

        btable = doc.add_table(rows=2, cols=4)
        btable.style     = 'Table Grid'
        btable.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["Hak Edilen", "Kullanılan", "Bekleyen", "Kalan"]):
            c = btable.cell(0, i)
            c.text = h
            for p in c.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(11)
        for i, v in enumerate([
                f"{entitled:.0f}", f"{used:.0f}",
                f"{pending:.0f}", f"{remaining:.0f}"]):
            btable.cell(1, i).text = v

        doc.add_paragraph()
        st = doc.add_table(rows=2, cols=2)
        st.autofit = False
        st.columns[0].width = st.columns[1].width = Inches(3)
        for (r_idx, c_idx, text) in [
            (0, 0, "Personel:\nAd Soyad: _________________\nİmza: _________________"),
            (0, 1, "İdari Kullanım\nKarar: ______________\nYönetici: ______________\nİmza: ______________"),
            (1, 0, "Teslim Tarihi: ___________"),
            (1, 1, "Onay: □ Evet   □ Hayır"),
        ]:
            p = st.cell(r_idx, c_idx).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(text)

        doc.add_paragraph()
        f = doc.add_paragraph(
            "Bu belge İnsan Kaynakları Yönetim Sistemi tarafından oluşturulmuştur.")
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in f.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)

        try:
            filename = (f"izin_talebi_{req_id}_"
                        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            filepath = os.path.join(os.path.expanduser("~"), "Downloads", filename)
            doc.save(filepath)
            os.startfile(filepath)
            QMessageBox.information(
                self, "Başarılı",
                f"Word dosyası oluşturuldu.\nKonum: {filepath}")
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Dosya kaydedilemedi:\n{e}")


# ============================================================
# نافذة إضافة/تعديل نوع الإجازة
# ============================================================
class EditLeaveTypeDialog(QDialog):
    def __init__(self, parent, db: DatabaseManager, user: dict,
                 comm=None, type_id=None):
        super().__init__(parent)
        self.db      = db
        self.user    = user
        self.comm    = comm
        self.type_id = type_id
        self.setWindowTitle("إضافة نوع إجازة" if not type_id else "تعديل نوع إجازة")
        self.setFixedSize(400, 350)
        self.setLayoutDirection(Qt.RightToLeft)
        self._build()
        if type_id:
            self._load()

    def _build(self):
        lay  = QVBoxLayout(self)
        form = QFormLayout()
        form.setSpacing(10)

        self.name_edit         = QLineEdit()
        self.days_spin         = QSpinBox()
        self.days_spin.setRange(0, 365)
        self.paid_combo        = QComboBox()
        self.paid_combo.addItems(["نعم", "لا"])
        self.carry_combo       = QComboBox()
        self.carry_combo.addItems(["نعم", "لا"])
        self.annual_combo      = QComboBox()
        self.annual_combo.addItems(["لا", "نعم"])
        self.max_requests_spin = QSpinBox()
        self.max_requests_spin.setRange(0, 100)
        self.max_requests_spin.setSpecialValueText("غير محدود")

        form.addRow("الاسم:",                              self.name_edit)
        form.addRow("أيام/سنة:",                          self.days_spin)
        form.addRow("مدفوعة:",                             self.paid_combo)
        form.addRow("ترحيل:",                              self.carry_combo)
        form.addRow("سنوية:",                              self.annual_combo)
        form.addRow("الحد الأقصى للطلبات (0=غير محدود):", self.max_requests_spin)
        lay.addLayout(form)

        btns = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        btns.accepted.connect(self._save)
        btns.rejected.connect(self.reject)
        lay.addWidget(btns)

    def _load(self):
        data = self.db.fetch_one(
            "SELECT name, days_per_year, paid, carry_over, is_annual, max_requests "
            "FROM leave_types WHERE id=?", (self.type_id,))
        if data:
            self.name_edit.setText(data[0])
            self.days_spin.setValue(data[1])
            self.paid_combo.setCurrentIndex(0 if data[2] else 1)
            self.carry_combo.setCurrentIndex(0 if data[3] else 1)
            self.annual_combo.setCurrentIndex(1 if data[4] else 0)
            self.max_requests_spin.setValue(data[5] or 0)

    def _save(self):
        name = self.name_edit.text().strip()
        if not name:
            QMessageBox.warning(self, "خطأ", "الاسم مطلوب")
            return
        days    = self.days_spin.value()
        paid    = 1 if self.paid_combo.currentText()   == "نعم" else 0
        carry   = 1 if self.carry_combo.currentText()  == "نعم" else 0
        annual  = 1 if self.annual_combo.currentText() == "نعم" else 0
        max_req = self.max_requests_spin.value() or None

        if self.type_id:
            old = self.db.fetch_one(
                "SELECT name, days_per_year FROM leave_types WHERE id=?", (self.type_id,))
            self.db.execute_query(
                """UPDATE leave_types
                   SET name=?, days_per_year=?, paid=?, carry_over=?,
                       is_annual=?, max_requests=?
                   WHERE id=?""",
                (name, days, paid, carry, annual, max_req, self.type_id))
            self.db.log_action("تعديل نوع إجازة", "leave_types", self.type_id,
                               {"name": old[0]} if old else None,
                               {"name": name, "days_per_year": days})
        else:
            self.db.execute_query(
                """INSERT INTO leave_types
                   (name, days_per_year, paid, carry_over, is_annual, max_requests)
                   VALUES (?,?,?,?,?,?)""",
                (name, days, paid, carry, annual, max_req))
            self.db.log_action("إضافة نوع إجازة", "leave_types", self.db.last_id(),
                               None, {"name": name, "days_per_year": days})
        self.accept()


# ============================================================
# نافذة إضافة/تعديل رصيد إجازة
# ============================================================
class AddBalanceDialog(QDialog):
    def __init__(self, parent, db: DatabaseManager, user: dict,
                 year: int, comm=None, record=None):
        super().__init__(parent)
        self.db     = db
        self.user   = user
        self.comm   = comm
        self.year   = year
        self.record = record
        self.setWindowTitle(
            "إضافة رصيد إجازة" if not record else "تعديل رصيد إجازة")
        self.setFixedSize(400, 300)
        self.setLayoutDirection(Qt.RightToLeft)
        self._build()
        if record:
            self._load()

    def _build(self):
        lay  = QVBoxLayout(self)
        form = QFormLayout()

        self.emp_combo = QComboBox()
        for eid, name in self.db.fetch_all(
                "SELECT id, first_name||' '||last_name "
                "FROM employees WHERE status='نشط'"):
            self.emp_combo.addItem(name, eid)
        self.emp_combo.setEditable(True)
        self.emp_combo.completer().setCompletionMode(QCompleter.PopupCompletion)
        self.emp_combo.completer().setFilterMode(Qt.MatchContains)

        self.type_combo = QComboBox()
        for tid, tname in self.db.fetch_all("SELECT id, name FROM leave_types"):
            self.type_combo.addItem(tname, tid)
        self.type_combo.setEditable(True)
        self.type_combo.completer().setCompletionMode(QCompleter.PopupCompletion)
        self.type_combo.completer().setFilterMode(Qt.MatchContains)

        self.total_days   = QDoubleSpinBox()
        self.used_days    = QDoubleSpinBox()
        self.pending_days = QDoubleSpinBox()
        for sp in (self.total_days, self.used_days, self.pending_days):
            sp.setRange(0, 365)
            sp.setDecimals(1)

        form.addRow("الموظف:",          self.emp_combo)
        form.addRow("نوع الإجازة:",     self.type_combo)
        form.addRow("إجمالي الأيام:",   self.total_days)
        form.addRow("الأيام المستخدمة:", self.used_days)
        form.addRow("الأيام المعلقة:",  self.pending_days)
        lay.addLayout(form)

        btns = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        btns.accepted.connect(self._save)
        btns.rejected.connect(self.reject)
        lay.addWidget(btns)

    def _load(self):
        self.emp_combo.setCurrentIndex(self.emp_combo.findData(self.record[1]))
        self.type_combo.setCurrentIndex(self.type_combo.findData(self.record[2]))
        self.total_days.setValue(self.record[4])
        self.used_days.setValue(self.record[5])
        self.pending_days.setValue(self.record[6])

    def _save(self):
        emp_id  = self.emp_combo.currentData()
        type_id = self.type_combo.currentData()
        total   = self.total_days.value()
        used    = self.used_days.value()
        pending = self.pending_days.value()

        if self.record:
            self.db.execute_query(
                """UPDATE leave_balance
                   SET total_days=?, used_days=?, pending_days=?
                   WHERE id=?""",
                (total, used, pending, self.record[0]))
            self.db.log_action("تعديل رصيد إجازة", "leave_balance", self.record[0],
                               {"total_days": self.record[4]},
                               {"total_days": total})
        else:
            self.db.execute_query(
                """INSERT INTO leave_balance
                   (employee_id, leave_type_id, year, total_days, used_days, pending_days)
                   VALUES (?,?,?,?,?,?)""",
                (emp_id, type_id, self.year, total, used, pending))
            self.db.log_action("إضافة رصيد إجازة", "leave_balance", self.db.last_id(),
                               None, {"employee_id": emp_id, "total_days": total})
        self.accept()